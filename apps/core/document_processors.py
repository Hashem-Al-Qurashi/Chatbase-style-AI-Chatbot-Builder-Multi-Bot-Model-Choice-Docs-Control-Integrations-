"""
Document processing factory with specialized processors for different file types.
Implements enterprise-grade text extraction with proper error handling and monitoring.
"""

import io
import re
import time
from abc import ABC, abstractmethod
from typing import Dict, List, Optional, Tuple, Any
from dataclasses import dataclass
from pathlib import Path
import structlog

# Document processing libraries
import PyPDF2
from docx import Document as DocxDocument
import requests
from bs4 import BeautifulSoup

from .exceptions import (
    DocumentProcessingError,
    TextExtractionError,
    UnsupportedFileTypeError
)

logger = structlog.get_logger()


@dataclass
class ProcessedDocument:
    """Result of document processing."""
    text_content: str
    metadata: Dict[str, Any]
    word_count: int
    char_count: int
    pages: Optional[int] = None
    processing_time_ms: int = 0
    quality_score: float = 1.0


@dataclass
class DocumentMetadata:
    """Structured document metadata."""
    title: Optional[str] = None
    author: Optional[str] = None
    creation_date: Optional[str] = None
    modification_date: Optional[str] = None
    subject: Optional[str] = None
    keywords: List[str] = None
    language: Optional[str] = None
    page_count: Optional[int] = None
    file_size: Optional[int] = None


class DocumentProcessor(ABC):
    """
    Abstract base class for document processors.
    
    Each document type (PDF, DOCX, TXT) has its own processor
    that implements the specific extraction logic.
    """
    
    def __init__(self, timeout_seconds: int = 30):
        """
        Initialize processor.
        
        Args:
            timeout_seconds: Maximum processing time before timeout
        """
        self.timeout_seconds = timeout_seconds
        self.logger = structlog.get_logger().bind(processor=self.__class__.__name__)
    
    @abstractmethod
    def extract_text(self, file_content: bytes, filename: str) -> ProcessedDocument:
        """
        Extract text content from document.
        
        Args:
            file_content: Raw file bytes
            filename: Original filename for context
            
        Returns:
            ProcessedDocument: Extracted text and metadata
            
        Raises:
            TextExtractionError: If extraction fails
        """
        pass
    
    @abstractmethod
    def supports_file_type(self, mime_type: str) -> bool:
        """
        Check if this processor supports the given MIME type.
        
        Args:
            mime_type: MIME type to check
            
        Returns:
            bool: True if supported
        """
        pass
    
    def _calculate_quality_score(self, text: str) -> float:
        """
        Calculate quality score for extracted text.
        
        Args:
            text: Extracted text content
            
        Returns:
            float: Quality score between 0.0 and 1.0
        """
        if not text.strip():
            return 0.0
        
        # Basic quality indicators
        char_count = len(text)
        word_count = len(text.split())
        
        # Check for reasonable text patterns
        alpha_ratio = sum(c.isalpha() for c in text) / len(text) if text else 0
        space_ratio = text.count(' ') / len(text) if text else 0
        
        # Penalize for too many special characters or garbled text
        special_char_ratio = sum(c in '!@#$%^&*()_+-=[]{}|;:,.<>?' for c in text) / len(text) if text else 0
        
        quality_score = 1.0
        
        # Adjust based on text characteristics
        if alpha_ratio < 0.3:  # Less than 30% alphabetic characters
            quality_score *= 0.5
        
        if special_char_ratio > 0.1:  # More than 10% special characters
            quality_score *= 0.7
        
        if word_count < 10:  # Very short documents
            quality_score *= 0.8
        
        return max(0.0, min(1.0, quality_score))
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text.
        
        Args:
            text: Raw extracted text
            
        Returns:
            str: Cleaned text
        """
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Normalize quotes
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        # Remove repeated punctuation
        text = re.sub(r'\.{3,}', '...', text)
        text = re.sub(r'-{3,}', '---', text)
        
        return text.strip()


class PDFProcessor(DocumentProcessor):
    """Processor for PDF documents using PyPDF2."""
    
    def supports_file_type(self, mime_type: str) -> bool:
        """Check if this processor supports PDF files."""
        return mime_type == 'application/pdf'
    
    def extract_text(self, file_content: bytes, filename: str) -> ProcessedDocument:
        """
        Extract text from PDF using PyPDF2.
        
        Args:
            file_content: PDF file bytes
            filename: Original filename
            
        Returns:
            ProcessedDocument: Extracted text and metadata
            
        Raises:
            TextExtractionError: If PDF processing fails
        """
        start_time = time.time()
        
        try:
            self.logger.info("Starting PDF text extraction", filename=filename)
            
            # Create file-like object from bytes
            pdf_file = io.BytesIO(file_content)
            
            # Read PDF
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if pdf_reader.is_encrypted:
                raise TextExtractionError("PDF is encrypted and cannot be processed")
            
            # Extract metadata
            metadata = self._extract_pdf_metadata(pdf_reader)
            
            # Extract text from all pages
            text_content = ""
            page_count = len(pdf_reader.pages)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content += f"\n--- Page {page_num + 1} ---\n"
                        text_content += page_text
                except Exception as e:
                    self.logger.warning(
                        "Failed to extract text from PDF page",
                        page_num=page_num + 1,
                        error=str(e)
                    )
                    continue
            
            # Clean and process text
            cleaned_text = self._clean_text(text_content)
            
            if not cleaned_text.strip():
                raise TextExtractionError("No readable text found in PDF")
            
            # Calculate metrics
            word_count = len(cleaned_text.split())
            char_count = len(cleaned_text)
            quality_score = self._calculate_quality_score(cleaned_text)
            processing_time_ms = int((time.time() - start_time) * 1000)
            
            self.logger.info(
                "PDF text extraction completed",
                filename=filename,
                pages=page_count,
                word_count=word_count,
                char_count=char_count,
                quality_score=quality_score,
                processing_time_ms=processing_time_ms
            )
            
            return ProcessedDocument(
                text_content=cleaned_text,
                metadata=metadata,
                word_count=word_count,
                char_count=char_count,
                pages=page_count,
                processing_time_ms=processing_time_ms,
                quality_score=quality_score
            )
            
        except PyPDF2.errors.PdfReadError as e:
            raise TextExtractionError(f"Failed to read PDF: {str(e)}")
        except Exception as e:
            self.logger.error(
                "PDF processing failed",
                filename=filename,
                error=str(e),
                error_type=type(e).__name__
            )
            raise TextExtractionError(f"PDF processing failed: {str(e)}")
    
    def _extract_pdf_metadata(self, pdf_reader: PyPDF2.PdfReader) -> Dict[str, Any]:
        """Extract metadata from PDF."""
        metadata = {}
        
        try:
            if hasattr(pdf_reader, 'metadata') and pdf_reader.metadata:
                pdf_meta = pdf_reader.metadata
                
                metadata.update({
                    'title': pdf_meta.get('/Title'),
                    'author': pdf_meta.get('/Author'),
                    'subject': pdf_meta.get('/Subject'),
                    'creator': pdf_meta.get('/Creator'),
                    'producer': pdf_meta.get('/Producer'),
                    'creation_date': str(pdf_meta.get('/CreationDate', '')),
                    'modification_date': str(pdf_meta.get('/ModDate', '')),
                    'page_count': len(pdf_reader.pages),
                })
        except Exception as e:
            self.logger.warning("Failed to extract PDF metadata", error=str(e))
        
        return metadata


class DOCXProcessor(DocumentProcessor):
    """Processor for DOCX documents using python-docx."""
    
    def supports_file_type(self, mime_type: str) -> bool:
        """Check if this processor supports DOCX files."""
        return mime_type in [
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword'
        ]
    
    def extract_text(self, file_content: bytes, filename: str) -> ProcessedDocument:
        """
        Extract text from DOCX using python-docx.
        
        Args:
            file_content: DOCX file bytes
            filename: Original filename
            
        Returns:
            ProcessedDocument: Extracted text and metadata
            
        Raises:
            TextExtractionError: If DOCX processing fails
        """
        start_time = time.time()
        
        try:
            self.logger.info("Starting DOCX text extraction", filename=filename)
            
            # Create file-like object from bytes
            docx_file = io.BytesIO(file_content)
            
            # Read DOCX
            doc = DocxDocument(docx_file)
            
            # Extract metadata
            metadata = self._extract_docx_metadata(doc)
            
            # Extract text from paragraphs
            text_content = ""
            paragraph_count = 0
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
                    paragraph_count += 1
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            # Clean and process text
            cleaned_text = self._clean_text(text_content)
            
            if not cleaned_text.strip():
                raise TextExtractionError("No readable text found in DOCX")
            
            # Calculate metrics
            word_count = len(cleaned_text.split())
            char_count = len(cleaned_text)
            quality_score = self._calculate_quality_score(cleaned_text)
            processing_time_ms = int((time.time() - start_time) * 1000)
            
            metadata['paragraph_count'] = paragraph_count
            metadata['table_count'] = len(doc.tables)
            
            self.logger.info(
                "DOCX text extraction completed",
                filename=filename,
                paragraphs=paragraph_count,
                tables=len(doc.tables),
                word_count=word_count,
                char_count=char_count,
                quality_score=quality_score,
                processing_time_ms=processing_time_ms
            )
            
            return ProcessedDocument(
                text_content=cleaned_text,
                metadata=metadata,
                word_count=word_count,
                char_count=char_count,
                processing_time_ms=processing_time_ms,
                quality_score=quality_score
            )
            
        except Exception as e:
            self.logger.error(
                "DOCX processing failed",
                filename=filename,
                error=str(e),
                error_type=type(e).__name__
            )
            raise TextExtractionError(f"DOCX processing failed: {str(e)}")
    
    def _extract_docx_metadata(self, doc: DocxDocument) -> Dict[str, Any]:
        """Extract metadata from DOCX."""
        metadata = {}
        
        try:
            core_props = doc.core_properties
            
            metadata.update({
                'title': core_props.title,
                'author': core_props.author,
                'subject': core_props.subject,
                'keywords': core_props.keywords,
                'comments': core_props.comments,
                'category': core_props.category,
                'created': str(core_props.created) if core_props.created else None,
                'modified': str(core_props.modified) if core_props.modified else None,
                'last_modified_by': core_props.last_modified_by,
                'revision': core_props.revision,
            })
        except Exception as e:
            self.logger.warning("Failed to extract DOCX metadata", error=str(e))
        
        return metadata


class TextProcessor(DocumentProcessor):
    """Processor for plain text documents."""
    
    def supports_file_type(self, mime_type: str) -> bool:
        """Check if this processor supports text files."""
        return mime_type in ['text/plain', 'text/markdown']
    
    def extract_text(self, file_content: bytes, filename: str) -> ProcessedDocument:
        """
        Extract text from plain text files.
        
        Args:
            file_content: Text file bytes
            filename: Original filename
            
        Returns:
            ProcessedDocument: Extracted text and metadata
            
        Raises:
            TextExtractionError: If text processing fails
        """
        start_time = time.time()
        
        try:
            self.logger.info("Starting text extraction", filename=filename)
            
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text_content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    text_content = file_content.decode(encoding)
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if text_content is None:
                raise TextExtractionError("Unable to decode text file with any supported encoding")
            
            # Clean and process text
            cleaned_text = self._clean_text(text_content)
            
            if not cleaned_text.strip():
                raise TextExtractionError("No readable text found in file")
            
            # Calculate metrics
            word_count = len(cleaned_text.split())
            char_count = len(cleaned_text)
            line_count = len(cleaned_text.splitlines())
            quality_score = self._calculate_quality_score(cleaned_text)
            processing_time_ms = int((time.time() - start_time) * 1000)
            
            metadata = {
                'encoding': used_encoding,
                'line_count': line_count,
                'file_size_bytes': len(file_content),
            }
            
            self.logger.info(
                "Text extraction completed",
                filename=filename,
                encoding=used_encoding,
                lines=line_count,
                word_count=word_count,
                char_count=char_count,
                quality_score=quality_score,
                processing_time_ms=processing_time_ms
            )
            
            return ProcessedDocument(
                text_content=cleaned_text,
                metadata=metadata,
                word_count=word_count,
                char_count=char_count,
                processing_time_ms=processing_time_ms,
                quality_score=quality_score
            )
            
        except Exception as e:
            self.logger.error(
                "Text processing failed",
                filename=filename,
                error=str(e),
                error_type=type(e).__name__
            )
            raise TextExtractionError(f"Text processing failed: {str(e)}")


class DocumentProcessorFactory:
    """
    Factory for creating document processors based on MIME type.
    
    Implements the Factory pattern for extensible document processing.
    """
    
    def __init__(self):
        """Initialize factory with available processors."""
        self._processors = [
            PDFProcessor(),
            DOCXProcessor(),
            TextProcessor(),
        ]
        self.logger = structlog.get_logger().bind(component="DocumentProcessorFactory")
    
    def create_processor(self, mime_type: str) -> DocumentProcessor:
        """
        Create appropriate processor for the given MIME type.
        
        Args:
            mime_type: MIME type of the document
            
        Returns:
            DocumentProcessor: Processor for the given type
            
        Raises:
            UnsupportedFileTypeError: If no processor supports the type
        """
        for processor in self._processors:
            if processor.supports_file_type(mime_type):
                self.logger.info(
                    "Created document processor",
                    mime_type=mime_type,
                    processor_type=type(processor).__name__
                )
                return processor
        
        self.logger.error(
            "No processor found for MIME type",
            mime_type=mime_type,
            available_processors=[type(p).__name__ for p in self._processors]
        )
        raise UnsupportedFileTypeError(mime_type)
    
    def get_supported_mime_types(self) -> List[str]:
        """
        Get list of all supported MIME types.
        
        Returns:
            List[str]: Supported MIME types
        """
        supported_types = []
        test_types = [
            'application/pdf',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            'application/msword',
            'text/plain',
            'text/markdown'
        ]
        
        for mime_type in test_types:
            for processor in self._processors:
                if processor.supports_file_type(mime_type):
                    supported_types.append(mime_type)
                    break
        
        return supported_types
    
    def register_processor(self, processor: DocumentProcessor):
        """
        Register a new processor.
        
        Args:
            processor: New processor to register
        """
        self._processors.append(processor)
        self.logger.info(
            "Registered new document processor",
            processor_type=type(processor).__name__
        )


# Global factory instance
document_processor_factory = DocumentProcessorFactory()


def process_document(file_content: bytes, filename: str, mime_type: str) -> ProcessedDocument:
    """
    Convenience function to process a document.
    
    Args:
        file_content: Raw file bytes
        filename: Original filename
        mime_type: MIME type of the file
        
    Returns:
        ProcessedDocument: Processed document with text and metadata
        
    Raises:
        UnsupportedFileTypeError: If file type is not supported
        TextExtractionError: If processing fails
    """
    processor = document_processor_factory.create_processor(mime_type)
    return processor.extract_text(file_content, filename)