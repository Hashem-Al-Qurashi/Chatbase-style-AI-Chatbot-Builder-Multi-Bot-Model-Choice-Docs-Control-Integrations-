"""
Document processing pipeline for RAG system.
Handles PDF, DOCX, URLs with privacy enforcement and chunking strategies.
"""

import io
import re
from abc import ABC, abstractmethod
from typing import Any, Dict, List, Optional, Tuple, Union
from dataclasses import dataclass
from enum import Enum
import logging
import hashlib

import requests
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import tiktoken

from django.core.cache import cache
from django.utils import timezone

from apps.core.interfaces import FileStorage
from chatbot_saas.config import get_settings


settings = get_settings()
logger = logging.getLogger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    TXT = "txt"
    HTML = "html"
    MARKDOWN = "markdown"
    URL = "url"


class PrivacyLevel(Enum):
    """Privacy levels for document content."""
    PRIVATE = "private"          # Not searchable, not citable, only for training
    CITABLE = "citable"         # Searchable and citable with attribution
    PUBLIC = "public"           # Fully searchable and citable


@dataclass
class DocumentContent:
    """Extracted document content with metadata."""
    text: str
    title: Optional[str]
    author: Optional[str]
    metadata: Dict[str, Any]
    privacy_level: PrivacyLevel
    source_type: DocumentType
    content_hash: str
    token_count: int


@dataclass
class DocumentChunk:
    """Document chunk with privacy enforcement."""
    content: str
    chunk_index: int
    start_char: int
    end_char: int
    token_count: int
    privacy_level: PrivacyLevel
    metadata: Dict[str, Any]
    content_hash: str


class DocumentProcessor(ABC):
    """Abstract base class for document processors."""
    
    @abstractmethod
    def can_process(self, content_type: str, file_extension: str) -> bool:
        """Check if processor can handle the document type."""
        pass
    
    @abstractmethod
    def extract_content(
        self,
        file_content: bytes,
        filename: str,
        privacy_level: PrivacyLevel
    ) -> DocumentContent:
        """Extract content from document."""
        pass


class PDFProcessor(DocumentProcessor):
    """PDF document processor."""
    
    def can_process(self, content_type: str, file_extension: str) -> bool:
        """Check if can process PDF."""
        return content_type == "application/pdf" or file_extension.lower() == ".pdf"
    
    def _clean_pdf_text(self, text: str) -> str:
        """Clean up common PDF extraction issues."""
        # Fix missing spaces between words (common in PDFs)
        # Look for lowercase letter followed by uppercase without space
        text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)
        
        # Fix missing spaces after periods
        text = re.sub(r'\.([A-Z])', r'. \1', text)
        
        # Fix ligature issues (fi, fl, etc.)
        ligatures = {
            'ﬁ': 'fi',
            'ﬂ': 'fl',
            'ﬀ': 'ff',
            'ﬃ': 'ffi',
            'ﬄ': 'ffl'
        }
        for lig, replacement in ligatures.items():
            text = text.replace(lig, replacement)
        
        # Remove excessive whitespace but preserve paragraph breaks
        lines = text.split('\n')
        cleaned_lines = []
        for line in lines:
            # Clean up each line
            line = ' '.join(line.split())  # Normalize whitespace
            if line:  # Only add non-empty lines
                cleaned_lines.append(line)
        
        # Rejoin with proper paragraph breaks
        text = '\n'.join(cleaned_lines)
        
        # Ensure proper spacing after punctuation
        text = re.sub(r'([.!?])([A-Za-z])', r'\1 \2', text)
        
        return text
    
    def extract_content(
        self,
        file_content: bytes,
        filename: str,
        privacy_level: PrivacyLevel
    ) -> DocumentContent:
        """Extract text from PDF."""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PdfReader(pdf_file)
            
            # Extract text from all pages with improved formatting
            text_parts = []
            for page_num, page in enumerate(pdf_reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    # Clean up common PDF extraction issues
                    text = self._clean_pdf_text(text)
                    # Add page separator for better chunking
                    text_parts.append(f"--- Page {page_num} ---\n{text}")
            
            full_text = "\n\n".join(text_parts)
            
            # Extract metadata
            metadata = {}
            if pdf_reader.metadata:
                metadata.update({
                    "title": pdf_reader.metadata.get("/Title", ""),
                    "author": pdf_reader.metadata.get("/Author", ""),
                    "subject": pdf_reader.metadata.get("/Subject", ""),
                    "creator": pdf_reader.metadata.get("/Creator", ""),
                    "producer": pdf_reader.metadata.get("/Producer", ""),
                    "creation_date": pdf_reader.metadata.get("/CreationDate", ""),
                    "modification_date": pdf_reader.metadata.get("/ModDate", "")
                })
            
            metadata.update({
                "page_count": len(pdf_reader.pages),
                "filename": filename,
                "processing_method": "PyPDF2"
            })
            
            # Calculate content hash
            content_hash = hashlib.sha256(full_text.encode()).hexdigest()
            
            # Count tokens
            encoding = tiktoken.get_encoding("cl100k_base")
            token_count = len(encoding.encode(full_text))
            
            return DocumentContent(
                text=full_text,
                title=metadata.get("title") or filename,
                author=metadata.get("author"),
                metadata=metadata,
                privacy_level=privacy_level,
                source_type=DocumentType.PDF,
                content_hash=content_hash,
                token_count=token_count
            )
            
        except Exception as e:
            logger.error(f"PDF processing failed for {filename}: {str(e)}")
            raise ValueError(f"Failed to process PDF: {str(e)}")


class DOCXProcessor(DocumentProcessor):
    """DOCX document processor."""
    
    def can_process(self, content_type: str, file_extension: str) -> bool:
        """Check if can process DOCX."""
        return (
            content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            or file_extension.lower() == ".docx"
        )
    
    def extract_content(
        self,
        file_content: bytes,
        filename: str,
        privacy_level: PrivacyLevel
    ) -> DocumentContent:
        """Extract text from DOCX."""
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            # Extract text from paragraphs
            text_parts = []
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_parts)
            
            # Extract metadata from core properties
            metadata = {
                "filename": filename,
                "processing_method": "python-docx"
            }
            
            if hasattr(doc, 'core_properties'):
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "subject": core_props.subject or "",
                    "description": core_props.comments or "",
                    "keywords": core_props.keywords or "",
                    "category": core_props.category or "",
                    "created": str(core_props.created) if core_props.created else "",
                    "modified": str(core_props.modified) if core_props.modified else "",
                    "last_modified_by": core_props.last_modified_by or ""
                })
            
            # Count paragraphs and tables
            metadata.update({
                "paragraph_count": len(doc.paragraphs),
                "table_count": len(doc.tables)
            })
            
            # Calculate content hash
            content_hash = hashlib.sha256(full_text.encode()).hexdigest()
            
            # Count tokens
            encoding = tiktoken.get_encoding("cl100k_base")
            token_count = len(encoding.encode(full_text))
            
            return DocumentContent(
                text=full_text,
                title=metadata.get("title") or filename,
                author=metadata.get("author"),
                metadata=metadata,
                privacy_level=privacy_level,
                source_type=DocumentType.DOCX,
                content_hash=content_hash,
                token_count=token_count
            )
            
        except Exception as e:
            logger.error(f"DOCX processing failed for {filename}: {str(e)}")
            raise ValueError(f"Failed to process DOCX: {str(e)}")


class TextProcessor(DocumentProcessor):
    """Plain text document processor."""
    
    def can_process(self, content_type: str, file_extension: str) -> bool:
        """Check if can process text files."""
        return (
            content_type in ["text/plain", "text/markdown"]
            or file_extension.lower() in [".txt", ".md", ".markdown"]
        )
    
    def extract_content(
        self,
        file_content: bytes,
        filename: str,
        privacy_level: PrivacyLevel
    ) -> DocumentContent:
        """Extract content from text file."""
        try:
            # Try different encodings
            encodings = ["utf-8", "utf-16", "latin-1", "cp1252"]
            full_text = None
            
            for encoding in encodings:
                try:
                    full_text = file_content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            
            if full_text is None:
                raise ValueError("Could not decode text file with any encoding")
            
            # Basic metadata
            metadata = {
                "filename": filename,
                "processing_method": "text_decoder",
                "character_count": len(full_text),
                "line_count": len(full_text.splitlines())
            }
            
            # Detect if it's markdown
            if filename.lower().endswith(('.md', '.markdown')):
                metadata["format"] = "markdown"
                # Extract title from first heading
                lines = full_text.splitlines()
                title = filename
                for line in lines:
                    if line.startswith('# '):
                        title = line[2:].strip()
                        break
                metadata["title"] = title
            
            # Calculate content hash
            content_hash = hashlib.sha256(full_text.encode()).hexdigest()
            
            # Count tokens
            encoding = tiktoken.get_encoding("cl100k_base")
            token_count = len(encoding.encode(full_text))
            
            return DocumentContent(
                text=full_text,
                title=metadata.get("title", filename),
                author=None,
                metadata=metadata,
                privacy_level=privacy_level,
                source_type=DocumentType.TXT if not filename.lower().endswith(('.md', '.markdown')) else DocumentType.MARKDOWN,
                content_hash=content_hash,
                token_count=token_count
            )
            
        except Exception as e:
            logger.error(f"Text processing failed for {filename}: {str(e)}")
            raise ValueError(f"Failed to process text file: {str(e)}")


class URLProcessor(DocumentProcessor):
    """URL content processor."""
    
    def can_process(self, content_type: str, file_extension: str) -> bool:
        """Check if can process URLs."""
        return content_type == "text/html" or file_extension == ""
    
    def extract_content(
        self,
        file_content: bytes,
        filename: str,  # This will be the URL
        privacy_level: PrivacyLevel
    ) -> DocumentContent:
        """Extract content from URL."""
        try:
            url = filename  # filename parameter contains the URL
            
            # Fetch content with timeout and user agent
            headers = {
                "User-Agent": "Mozilla/5.0 (compatible; ChatbotSaaS/1.0; +https://example.com/bot)"
            }
            
            response = requests.get(url, headers=headers, timeout=30)
            response.raise_for_status()
            
            # Parse HTML content
            soup = BeautifulSoup(response.content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Extract title
            title_element = soup.find('title')
            title = title_element.get_text().strip() if title_element else url
            
            # Extract main content
            # Look for common content containers
            content_selectors = [
                'main', 'article', '.content', '#content',
                '.post', '.entry', '.article-content'
            ]
            
            main_content = None
            for selector in content_selectors:
                elements = soup.select(selector)
                if elements:
                    main_content = elements[0]
                    break
            
            # If no main content found, use body
            if main_content is None:
                main_content = soup.find('body') or soup
            
            # Extract text
            text = main_content.get_text()
            
            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            full_text = '\n'.join(chunk for chunk in chunks if chunk)
            
            # Extract metadata
            metadata = {
                "url": url,
                "title": title,
                "processing_method": "beautifulsoup",
                "content_type": response.headers.get('content-type', ''),
                "status_code": response.status_code,
                "extracted_at": timezone.now().isoformat()
            }
            
            # Try to extract meta description
            meta_desc = soup.find('meta', attrs={'name': 'description'})
            if meta_desc:
                metadata["description"] = meta_desc.get('content', '')
            
            # Try to extract author
            author_meta = soup.find('meta', attrs={'name': 'author'})
            if author_meta:
                metadata["author"] = author_meta.get('content', '')
            
            # Calculate content hash
            content_hash = hashlib.sha256(full_text.encode()).hexdigest()
            
            # Count tokens
            encoding = tiktoken.get_encoding("cl100k_base")
            token_count = len(encoding.encode(full_text))
            
            return DocumentContent(
                text=full_text,
                title=title,
                author=metadata.get("author"),
                metadata=metadata,
                privacy_level=privacy_level,
                source_type=DocumentType.URL,
                content_hash=content_hash,
                token_count=token_count
            )
            
        except Exception as e:
            logger.error(f"URL processing failed for {filename}: {str(e)}")
            raise ValueError(f"Failed to process URL: {str(e)}")


class DocumentProcessingPipeline:
    """Main document processing pipeline."""
    
    def __init__(self):
        self.processors = [
            PDFProcessor(),
            DOCXProcessor(),
            TextProcessor(),
            URLProcessor()
        ]
        self.cache_ttl = settings.CACHE_TTL_SECONDS
    
    def process_document(
        self,
        file_content: bytes,
        filename: str,
        content_type: str,
        privacy_level: PrivacyLevel,
        use_cache: bool = True
    ) -> DocumentContent:
        """
        Process document and extract content.
        
        Args:
            file_content: Raw file content
            filename: Original filename or URL
            content_type: MIME type
            privacy_level: Privacy level for content
            use_cache: Whether to use caching
            
        Returns:
            DocumentContent: Extracted content with metadata
        """
        # Generate cache key based on content hash
        content_hash = hashlib.sha256(file_content).hexdigest()
        cache_key = f"doc_content:{content_hash}:{privacy_level.value}"
        
        # Try cache first
        if use_cache and settings.ENABLE_CACHING:
            cached_content = cache.get(cache_key)
            if cached_content:
                logger.info(f"Using cached content for {filename}")
                return cached_content
        
        # Find appropriate processor
        file_extension = self._get_file_extension(filename)
        processor = None
        
        for proc in self.processors:
            if proc.can_process(content_type, file_extension):
                processor = proc
                break
        
        if not processor:
            raise ValueError(f"No processor found for content type: {content_type}")
        
        # Process document
        logger.info(f"Processing document: {filename} with {processor.__class__.__name__}")
        content = processor.extract_content(file_content, filename, privacy_level)
        
        # Validate content
        if not content.text or content.text.strip() == "":
            raise ValueError("No text content extracted from document")
        
        if content.token_count == 0:
            raise ValueError("Document appears to be empty")
        
        # Cache result
        if use_cache and settings.ENABLE_CACHING:
            cache.set(cache_key, content, timeout=self.cache_ttl)
        
        logger.info(f"Document processed successfully: {filename} ({content.token_count} tokens)")
        return content
    
    def process_url(
        self,
        url: str,
        privacy_level: PrivacyLevel,
        use_cache: bool = True
    ) -> DocumentContent:
        """
        Process URL and extract content.
        
        Args:
            url: URL to process
            privacy_level: Privacy level for content
            use_cache: Whether to use caching
            
        Returns:
            DocumentContent: Extracted content with metadata
        """
        # Generate cache key for URL
        url_hash = hashlib.sha256(url.encode()).hexdigest()
        cache_key = f"url_content:{url_hash}:{privacy_level.value}"
        
        # Try cache first
        if use_cache and settings.ENABLE_CACHING:
            cached_content = cache.get(cache_key)
            if cached_content:
                logger.info(f"Using cached content for URL: {url}")
                return cached_content
        
        # Use URL processor
        processor = URLProcessor()
        content = processor.extract_content(b"", url, privacy_level)
        
        # Cache result
        if use_cache and settings.ENABLE_CACHING:
            cache.set(cache_key, content, timeout=self.cache_ttl)
        
        logger.info(f"URL processed successfully: {url} ({content.token_count} tokens)")
        return content
    
    def _get_file_extension(self, filename: str) -> str:
        """Extract file extension from filename."""
        import os
        return os.path.splitext(filename)[1].lower()


# Global document processing pipeline
document_pipeline = DocumentProcessingPipeline()