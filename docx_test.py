#!/usr/bin/env python3
"""Create a DOCX test file to validate document processing."""

from docx import Document
import os

def create_test_docx():
    doc = Document()
    doc.add_heading('DOCX Content Validation Test', 0)
    
    para1 = doc.add_paragraph('This is a test DOCX document to validate the file content corruption fix. ')
    para1.add_run('Important validation phrase: ').bold = True
    para1.add_run('DOCX_PROCESSING_SUCCESS_VALIDATION')
    
    doc.add_paragraph('The file content should be extracted correctly from this DOCX format.')
    doc.add_paragraph('Test ID: GRUMPY_TESTER_DOCX_001')
    
    # Save in current directory
    docx_path = '/home/sakr_quraish/Projects/Ismail/test_document.docx'
    doc.save(docx_path)
    print(f"Created test DOCX: {docx_path}")
    return docx_path

if __name__ == "__main__":
    create_test_docx()